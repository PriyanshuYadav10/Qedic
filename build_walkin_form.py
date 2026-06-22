"""Generate a styled Walk-In Interview Candidate Registration Form (.docx)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


PRIMARY = RGBColor(0x1F, 0x4E, 0x79)   # deep blue
ACCENT = RGBColor(0x2E, 0x75, 0xB6)    # mid blue
LIGHT_BG = "DEEBF7"                    # light blue cell shade (hex string for XML)
SECTION_BG = "1F4E79"                  # section header fill
DIVIDER = RGBColor(0xBF, 0xBF, 0xBF)
TEXT = RGBColor(0x33, 0x33, 0x33)

CHECKBOX = "☐"   # ballot box


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFBFBF", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_table_borders(table, color="BFBFBF", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def add_horizontal_rule(paragraph, color="BFBFBF", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section_header(doc, text):
    """Full-width blue band heading."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, SECTION_BG)
    # remove all borders
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tc_borders.append(b)
    tc_pr.append(tc_borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # spacer after band
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_field(doc, label, width_label=Cm(4.5), width_value=Cm(12)):
    """One-row field: bold label | underlined blank value."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    label_cell, value_cell = table.rows[0].cells
    label_cell.width = width_label
    value_cell.width = width_value

    # label
    p = label_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = TEXT

    # value (blank with bottom border)
    p = value_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_horizontal_rule(p, color="7F7F7F", size="6")

    # strip cell borders
    for c in (label_cell, value_cell):
        tc_pr = c._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tc_borders.append(b)
        tc_pr.append(tc_borders)

    # spacing after
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.space_before = Pt(0)


def add_checkbox_row(doc, label, options, columns=None):
    """Row with bold label followed by checkbox options."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label + "  ")
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = TEXT

    for opt in options:
        r = p.add_run(f"{CHECKBOX} {opt}    ")
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT


def add_checkbox_list(doc, options):
    """Vertical list of checkboxes."""
    for opt in options:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f"{CHECKBOX}  {opt}")
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT


def add_blank_lines(doc, count=3):
    """Underlined blank lines for free-text answers."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_horizontal_rule(p, color="7F7F7F", size="6")


def style_header_cell(cell, text):
    set_cell_shading(cell, SECTION_BG)
    set_cell_borders(cell, color="1F4E79")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def style_data_cell(cell, text=""):
    set_cell_borders(cell, color="BFBFBF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if text:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = TEXT


def build():
    doc = Document()

    # page setup
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = TEXT

    # ── HEADER ─────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run("GrainZap IT Solutions Pvt. Ltd.")
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run("Walk-In Interview Candidate Registration Form")
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = ACCENT

    addr = doc.add_paragraph()
    addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    addr.paragraph_format.space_after = Pt(0)
    r = addr.add_run("Jaipur, Rajasthan  |  hr@grainzap.com  |  www.grainzap.com")
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    divider = doc.add_paragraph()
    add_horizontal_rule(divider, color="1F4E79", size="12")
    divider.paragraph_format.space_after = Pt(6)

    # ── 1. PERSONAL DETAILS ────────────────────────────────────────
    add_section_header(doc, "1. Personal Details")
    add_field(doc, "Full Name")
    add_field(doc, "Mobile Number")
    add_field(doc, "Email Address")
    add_field(doc, "Current City")
    add_field(doc, "Date of Birth")
    add_checkbox_row(doc, "Gender:", ["Male", "Female", "Other"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 2. POSITION APPLIED FOR ────────────────────────────────────
    add_section_header(doc, "2. Position Applied For")
    add_checkbox_list(doc, [
        "Flutter Developer",
        "Full Stack Developer",
        "Business Development Executive",
        "Digital Marketing Executive",
        "UI/UX Designer",
    ])
    add_field(doc, "Other (please specify)")

    # ── 3. EDUCATION ───────────────────────────────────────────────
    add_section_header(doc, "3. Educational Qualification")
    edu_table = doc.add_table(rows=5, cols=4)
    edu_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(edu_table, color="BFBFBF", size="6")
    headers = ["Qualification", "Institute / College", "Year of Passing", "Percentage / CGPA"]
    for i, h in enumerate(headers):
        style_header_cell(edu_table.rows[0].cells[i], h)
    rows_labels = ["10th", "12th", "Graduation", "Post-Graduation / Other"]
    for i, lbl in enumerate(rows_labels, start=1):
        for j in range(4):
            cell = edu_table.rows[i].cells[j]
            if j == 0:
                set_cell_shading(cell, LIGHT_BG)
                style_data_cell(cell, lbl)
            else:
                style_data_cell(cell)
            cell.height = Cm(0.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 4. TECHNICAL SKILLS ────────────────────────────────────────
    add_section_header(doc, "4. Technical Skills")
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(2)
    r = note.add_run("List programming languages, frameworks, tools, design software, etc.")
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    add_blank_lines(doc, count=4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 5. WORK EXPERIENCE ─────────────────────────────────────────
    add_section_header(doc, "5. Work Experience")
    add_checkbox_row(doc, "Status:", ["Fresher", "Experienced"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(2)
    r = note.add_run("If experienced, fill in the details below:")
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    add_field(doc, "Company Name")
    add_field(doc, "Role / Designation")
    add_field(doc, "Experience Duration")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Key Responsibilities:")
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.bold = True
    add_blank_lines(doc, count=2)

    # ── 6. INTERNSHIP / PROJECT DETAILS ───────────────────────────
    add_section_header(doc, "6. Internship / Project Details")
    add_field(doc, "Project Title")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Description:")
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.bold = True
    add_blank_lines(doc, count=3)
    add_field(doc, "Technologies Used")
    add_field(doc, "Duration")

    # ── 7. AVAILABILITY ────────────────────────────────────────────
    add_section_header(doc, "7. Availability Details")
    add_checkbox_row(doc, "Can you join immediately?", ["Yes", "No"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_field(doc, "Notice Period (if any)")

    # ── 8. ADDITIONAL INFORMATION ──────────────────────────────────
    add_section_header(doc, "8. Additional Information")
    add_checkbox_row(doc, "Do you have your own laptop?", ["Yes", "No"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_checkbox_row(doc, "Comfortable working from office in Jaipur?", ["Yes", "No"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_field(doc, "Expected Salary / Stipend")
    add_checkbox_row(
        doc,
        "How did you hear about us?",
        ["LinkedIn", "Naukri", "Referral", "Social Media", "Other"],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 9. DOCUMENT CHECKLIST ──────────────────────────────────────
    add_section_header(doc, "9. Document Checklist (to be brought)")
    add_checkbox_list(doc, [
        "Updated Resume / CV",
        "Aadhaar Card (original + photocopy)",
        "Passport Size Photograph (2 copies)",
        "Educational Documents (marksheets & certificates)",
        "Experience / Relieving Letter (if applicable)",
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 10. DECLARATION ────────────────────────────────────────────
    add_section_header(doc, "10. Declaration")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(
        "I hereby declare that all the information provided above is true to the best of my "
        "knowledge. I understand that any misrepresentation may lead to rejection of my "
        "candidature or termination of employment."
    )
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = TEXT

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    for i, (label, w) in enumerate([("Candidate Signature", Cm(9.5)), ("Date", Cm(6.5))]):
        cell = sig_table.rows[0].cells[i]
        cell.width = w
        # strip borders
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tc_borders.append(b)
        tc_pr.append(tc_borders)

        line_p = cell.paragraphs[0]
        line_p.paragraph_format.space_after = Pt(0)
        add_horizontal_rule(line_p, color="7F7F7F", size="6")
        lbl_p = cell.add_paragraph()
        lbl_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lbl_p.paragraph_format.space_after = Pt(0)
        run = lbl_p.add_run(label)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── OFFICE USE ONLY ────────────────────────────────────────────
    add_section_header(doc, "Office Use Only")
    office_table = doc.add_table(rows=5, cols=2)
    office_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(office_table, color="BFBFBF", size="6")
    office_rows = [
        ("Interview Round", ""),
        ("Interviewer Name", ""),
        ("Date of Interview", ""),
        ("Technical Rating (out of 10)", ""),
        ("Communication Rating (out of 10)", ""),
    ]
    for i, (label, val) in enumerate(office_rows):
        lcell = office_table.rows[i].cells[0]
        vcell = office_table.rows[i].cells[1]
        set_cell_shading(lcell, LIGHT_BG)
        style_data_cell(lcell, label)
        style_data_cell(vcell, val)
        lcell.width = Cm(7)
        vcell.width = Cm(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_checkbox_row(
        doc,
        "Status:",
        ["Selected", "On Hold", "Rejected", "Shortlisted for Next Round"],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Remarks:")
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.bold = True
    add_blank_lines(doc, count=2)

    add_field(doc, "Authorised Signature")

    out_path = "/Users/priyanshu/projects/Qedic/WalkInInterviewForm.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    build()
